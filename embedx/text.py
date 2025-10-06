import torch
import h5py
import os
import numpy as np
from tqdm.auto import tqdm
import fitz
from sentence_transformers import SentenceTransformer
import sys
import docx
from bs4 import BeautifulSoup

class EmbedX:
    def __init__(self):
        # Automatically select device: use GPU (CUDA) if available, otherwise CPU
        self.model = SentenceTransformer('all-roberta-large-v1')

        # Paths and storage for image dataset and image embeddings
        self.dataset_path = ""
        self.embedding_vector_path = ""
        self.vector = self.model.encode("", show_progress_bar=False, convert_to_numpy=True)

    def select_Dataset_Path(self, dataset_path):
        self.dataset_path = dataset_path

    def select_Output_Path(self, embedding_vector_path):
        self.embedding_vector_path = embedding_vector_path
        if not os.path.exists(embedding_vector_path):
            with h5py.File(self.embedding_vector_path, "w") as outfile:
                outfile.create_dataset(
                    "path",
                    shape=(0,),
                    maxshape=(None,),
                    dtype=h5py.string_dtype(encoding="utf-8")
                )
                
                outfile.create_dataset(
                    "embeddings",
                    shape=(0, 1024),
                    maxshape=(None, 1024),
                    dtype=np.float32,
                    chunks=True
                )

    def embed_Dataset(self):
        for filename in tqdm(os.listdir(self.dataset_path), desc="Embedding"):
            file_path = os.path.join(self.dataset_path, filename)
            text = ""
            if filename.endswith(".pdf"):
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text()
                doc.close()

            elif filename.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                text = text.strip()

            elif filename.endswith(".docx"):
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + " "
                text = text.strip()

            elif filename.endswith(".html") or filename.endswith(".htm"):
                with open(file_path, "r", encoding="utf-8") as f:
                    html_content = f.read()
                soup = BeautifulSoup(html_content, "lxml")
                text = soup.get_text(separator=" ", strip=True)
                text = text.strip()
            
            self.vector = self.model.encode(text, show_progress_bar=False, convert_to_numpy=True)
            self.vector = self.vector / np.linalg.norm(self.vector)
            with h5py.File(self.embedding_vector_path, "a") as f:
                    idx = f["path"].shape[0]

                    f["path"].resize(idx + 1, axis=0)
                    f["embeddings"].resize(idx + 1, axis=0)

                    f["path"][idx] = file_path
                    f["embeddings"][idx] = self.vector

    def embed_Other_file(self, file_path):
        pbar = tqdm(total = 1, desc = "Embedding: ")
        filename = os.path.basename(file_path)
        text = ""
        if filename.endswith(".pdf"):
            file_path = os.path.join(self.dataset_path, filename)
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()

        elif filename.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            text = text.strip()

        elif filename.endswith(".docx"):
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + " "
            text = text.strip()

        elif filename.endswith(".html") or filename.endswith(".htm"):
            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            soup = BeautifulSoup(html_content, "lxml")
            text = soup.get_text(separator=" ", strip=True)
            text = text.strip()
        
        self.vector = self.model.encode(text, show_progress_bar=False, convert_to_numpy=True)
        self.vector = self.vector / np.linalg.norm(self.vector)

        with h5py.File(self.embedding_vector_path, "a") as f:
            idx = f["path"].shape[0]

            f["path"].resize(idx + 1, axis=0)
            f["embeddings"].resize(idx + 1, axis=0)

            f["path"][idx] = file_path
            f["embeddings"][idx] = self.vector
        pbar.update(1)
        
def embed_Text(text):
    model = SentenceTransformer('all-roberta-large-v1')
    vector = model.encode(text, show_progress_bar=False, convert_to_numpy=True)
    vector = vector / np.linalg.norm(vector)
    return vector